"""Document parser — PDF and DOCX text extraction.

Uses pdfplumber for PDF files and python-docx for DOCX files.
Returns structured extraction results with page/section metadata.

Usage:
    result = parse_document(Path("resume.pdf"))
    print(result.full_text)
    for page in result.pages:
        print(page.page_number, page.text[:100])
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from logging_config import get_logger

logger = get_logger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx"}


@dataclass
class PageContent:
    """Extracted content from a single page or section.

    Attributes:
        page_number: 1-indexed page/section number.
        text: Extracted text content.
        char_count: Number of characters in the text.
    """

    page_number: int
    text: str
    char_count: int = 0

    def __post_init__(self) -> None:
        """Compute char_count after initialization."""
        self.char_count = len(self.text)


@dataclass
class ParseResult:
    """Result of parsing a single document.

    Attributes:
        file_name: Original file name.
        file_type: File extension (e.g., ".pdf", ".docx").
        pages: List of extracted page contents.
        full_text: Concatenated text from all pages.
        total_pages: Number of pages/sections extracted.
        metadata: Additional extraction metadata.
    """

    file_name: str
    file_type: str
    pages: list[PageContent] = field(default_factory=list)
    full_text: str = ""
    total_pages: int = 0
    metadata: dict[str, str | int] = field(default_factory=dict)


def _parse_pdf(file_path: Path) -> ParseResult:
    """Extract text from a PDF file using pdfplumber.

    Args:
        file_path: Path to the PDF file.

    Returns:
        ParseResult with per-page content and concatenated full text.

    Raises:
        ValueError: If the PDF contains no extractable text.
    """
    logger.info("parser.pdf.start", file=str(file_path))
    pages: list[PageContent] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw_text = page.extract_text() or ""
            # Normalize whitespace: collapse multiple spaces, strip lines
            cleaned = "\n".join(
                line.strip() for line in raw_text.splitlines() if line.strip()
            )
            pages.append(PageContent(page_number=i, text=cleaned))

        metadata: dict[str, str | int] = {
            "page_count": len(pdf.pages),
        }
        # Extract PDF metadata if available
        if pdf.metadata:
            for key in ("Title", "Author", "Creator"):
                if pdf.metadata.get(key):
                    metadata[key.lower()] = str(pdf.metadata[key])

    full_text = "\n\n".join(p.text for p in pages if p.text)

    if not full_text.strip():
        raise ValueError(f"No extractable text found in PDF: {file_path.name}")

    result = ParseResult(
        file_name=file_path.name,
        file_type=".pdf",
        pages=pages,
        full_text=full_text,
        total_pages=len(pages),
        metadata=metadata,
    )

    logger.info(
        "parser.pdf.complete",
        file=file_path.name,
        pages=len(pages),
        chars=len(full_text),
    )
    return result


def _parse_docx(file_path: Path) -> ParseResult:
    """Extract text from a DOCX file using python-docx.

    Extracts paragraph text and table content. Each logical section
    (group of paragraphs between headings) becomes a "page" in the result.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        ParseResult with section-based content.

    Raises:
        ValueError: If the DOCX contains no extractable text.
    """
    logger.info("parser.docx.start", file=str(file_path))
    doc = DocxDocument(str(file_path))

    # Extract all paragraph text
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract table content
    table_texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    # Combine paragraphs and tables
    all_text_parts = paragraphs + table_texts
    full_text = "\n".join(all_text_parts)

    if not full_text.strip():
        raise ValueError(f"No extractable text found in DOCX: {file_path.name}")

    # Group into sections based on headings or fixed-size groups
    sections: list[PageContent] = []
    current_section: list[str] = []
    section_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings — start a new section
        is_heading = (
            para.style
            and para.style.name
            and para.style.name.lower().startswith("heading")
        )

        if is_heading and current_section:
            sections.append(
                PageContent(page_number=section_num, text="\n".join(current_section))
            )
            section_num += 1
            current_section = [text]
        else:
            current_section.append(text)

    # Flush the last section
    if current_section:
        sections.append(
            PageContent(page_number=section_num, text="\n".join(current_section))
        )

    # If no heading-based sections were found, treat the whole doc as one section
    if not sections:
        sections = [PageContent(page_number=1, text=full_text)]

    # Append table content as a final section if present
    if table_texts:
        sections.append(
            PageContent(page_number=len(sections) + 1, text="\n".join(table_texts))
        )

    result = ParseResult(
        file_name=file_path.name,
        file_type=".docx",
        pages=sections,
        full_text=full_text,
        total_pages=len(sections),
        metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables)},
    )

    logger.info(
        "parser.docx.complete",
        file=file_path.name,
        sections=len(sections),
        chars=len(full_text),
    )
    return result


def parse_document(file_path: Path) -> ParseResult:
    """Parse a resume document (PDF or DOCX) and extract text.

    This is the main entry point for document parsing. It dispatches
    to the appropriate parser based on file extension.

    Args:
        file_path: Path to the document file.

    Returns:
        ParseResult containing extracted text with structural metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is unsupported or contains no text.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{extension}'. "
            f"Supported types: {SUPPORTED_EXTENSIONS}"
        )

    if extension == ".pdf":
        return _parse_pdf(file_path)
    else:
        return _parse_docx(file_path)
