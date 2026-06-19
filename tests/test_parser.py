"""Tests for preprocessing.parser module.

Tests cover PDF parsing, DOCX parsing, error handling,
and the unified parse_document entry point.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pdfplumber
import pytest
from docx import Document as DocxDocument

from preprocessing.parser import ParseResult, parse_document


# ── Fixtures: Create test files ────────────────────────────────────


def _create_test_pdf(path: Path, pages_text: list[str]) -> Path:
    """Create a minimal test PDF using pdfplumber's companion library.

    We use reportlab if available, otherwise create a simple PDF manually.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(str(path), pagesize=letter)
        for text in pages_text:
            # Write text line by line
            y = 750
            for line in text.splitlines():
                c.drawString(72, y, line)
                y -= 15
            c.showPage()
        c.save()
    except ImportError:
        # Fallback: create a minimal valid PDF manually
        # This is a bare-minimum PDF with text
        content = pages_text[0] if pages_text else "Test content"
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"4 0 obj<</Length " + str(len(content) + 30).encode() + b">>\n"
            b"stream\nBT /F1 12 Tf 72 720 Td (" + content.encode() + b") Tj ET\n"
            b"endstream\nendobj\n"
            b"xref\n0 6\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"0000000115 00000 n \n"
            b"0000000306 00000 n \n"
            b"0000000266 00000 n \n"
            b"trailer<</Size 6/Root 1 0 R>>\n"
            b"startxref\n0\n%%EOF"
        )
        path.write_bytes(pdf_bytes)
    return path


def _create_test_docx(path: Path, paragraphs: list[tuple[str, str]]) -> Path:
    """Create a test DOCX file with paragraphs.

    Args:
        path: Output file path.
        paragraphs: List of (text, style) tuples. Style can be "Normal" or "Heading 1" etc.
    """
    doc = DocxDocument()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style)
    doc.save(str(path))
    return path


# ── Tests ──────────────────────────────────────────────────────────


class TestParseDocument:
    """Tests for the parse_document entry point."""

    def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        """Unsupported file types should raise ValueError."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text("some text")
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_document(txt_file)

    def test_missing_file_raises(self, tmp_path: Path) -> None:
        """Non-existent files should raise FileNotFoundError."""
        missing = tmp_path / "nonexistent.pdf"
        with pytest.raises(FileNotFoundError):
            parse_document(missing)


class TestDocxParser:
    """Tests for DOCX parsing."""

    def test_basic_docx_parsing(self, tmp_path: Path) -> None:
        """Should extract text from a simple DOCX."""
        docx_path = tmp_path / "resume.docx"
        _create_test_docx(
            docx_path,
            [
                ("John Doe", "Normal"),
                ("Software Engineer", "Normal"),
                ("Experience", "Heading 1"),
                ("5 years at Google working on search infrastructure", "Normal"),
                ("Education", "Heading 1"),
                ("BS Computer Science, Stanford University", "Normal"),
            ],
        )

        result = parse_document(docx_path)

        assert isinstance(result, ParseResult)
        assert result.file_type == ".docx"
        assert result.file_name == "resume.docx"
        assert "John Doe" in result.full_text
        assert "Google" in result.full_text
        assert result.total_pages > 0

    def test_docx_sections_detected(self, tmp_path: Path) -> None:
        """Should detect heading-based sections in DOCX."""
        docx_path = tmp_path / "sections.docx"
        _create_test_docx(
            docx_path,
            [
                ("Summary Section", "Heading 1"),
                ("An experienced engineer.", "Normal"),
                ("Work Details", "Heading 1"),
                ("Worked at multiple companies.", "Normal"),
            ],
        )

        result = parse_document(docx_path)

        # Should have multiple sections based on headings
        assert result.total_pages >= 2

    def test_empty_docx_raises(self, tmp_path: Path) -> None:
        """Empty DOCX should raise ValueError."""
        docx_path = tmp_path / "empty.docx"
        doc = DocxDocument()
        doc.save(str(docx_path))

        with pytest.raises(ValueError, match="No extractable text"):
            parse_document(docx_path)

    def test_docx_with_tables(self, tmp_path: Path) -> None:
        """Should extract text from DOCX tables."""
        docx_path = tmp_path / "tables.docx"
        doc = DocxDocument()
        doc.add_paragraph("Resume Header", style="Normal")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"
        doc.save(str(docx_path))

        result = parse_document(docx_path)

        assert "Python" in result.full_text
        assert result.metadata.get("table_count") == 1

    def test_parse_result_metadata(self, tmp_path: Path) -> None:
        """ParseResult should include paragraph and table counts."""
        docx_path = tmp_path / "meta.docx"
        _create_test_docx(
            docx_path,
            [
                ("Line 1", "Normal"),
                ("Line 2", "Normal"),
                ("Line 3", "Normal"),
            ],
        )

        result = parse_document(docx_path)

        assert result.metadata["paragraph_count"] == 3
        assert result.metadata["table_count"] == 0
